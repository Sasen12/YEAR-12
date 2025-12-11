import io
from docx import Document
from fastapi.testclient import TestClient
from app.main import app

client = TestClient(app)

def make_docx_bytes(q_and_answers):
    doc = Document()
    for q, answers in q_and_answers:
        # question paragraph then answer paragraphs
        doc.add_paragraph(q)
        for a in answers:
            doc.add_paragraph(a)
        doc.add_paragraph('')
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def test_docx_import_and_perfect_score():
    # register and login
    client.post('/auth/register', json={'username':'docxuser','password':'pw'})
    r = client.post('/auth/login', json={'username':'docxuser','password':'pw'})
    assert r.status_code == 200
    token = r.json()['access_token']
    headers = {'Authorization': f'Bearer {token}'}

    qlist = [("What is 2+2?", ["4", "3"]), ("Capital of France?", ["Paris", "London"])]
    docx_bytes = make_docx_bytes(qlist)
    files = {'file': ('questions.docx', docx_bytes)}
    r2 = client.post('/softwaredev/import', files=files, headers=headers)
    assert r2.status_code == 200
    resjson = r2.json()
    assert resjson['created'] >= 2

    # fetch all questions and pick the ones we just imported by matching text
    all_q = client.get('/softwaredev/questions').json()
    answers = []
    for expected_q, _ in qlist:
        match = next((q for q in all_q if expected_q in q['question_text']), None)
        assert match is not None, f"Imported question not found: {expected_q}"
        given = match['answers'][0]['answer_text'] if match['answers'] else ''
        answers.append({'question_id': match['id'], 'given_answer': given})
    payload = {'answers': answers}
    g = client.post('/softwaredev/grade', json=payload, headers=headers)
    assert g.status_code == 200
    gj = g.json()
    assert gj['score'] == gj['total'] or gj['total'] == 0

def test_invalid_token_rejected():
    headers = {'Authorization': 'Bearer invalid.token.here'}
    r = client.get('/softwaredev/quiz', headers=headers)
    assert r.status_code == 401 or r.status_code == 403
